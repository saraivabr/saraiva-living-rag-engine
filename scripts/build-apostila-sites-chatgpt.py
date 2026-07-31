from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "artifacts"
OUTPUT_DOCX = OUTPUT_DIR / "apostila-sites-chatgpt.docx"
PROMPT_SOURCE = ROOT / "content" / "prompt-site-work-sites.md"

BLUE = "1666D3"
NAVY = "0B2545"
INK = "172033"
MUTED = "56657A"
LIGHT_BLUE = "EAF2FF"
PALE = "F5F8FC"
WHITE = "FFFFFF"
GREEN = "147D64"
RED = "9B1C1C"
GOLD = "B77700"
NEXT_NUM_ID = 100


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_numbering(doc):
    numbering = doc.part.numbering_part.element

    def create_abstract(abstract_id, fmt, text, align, left, hanging, font="Calibri"):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), align)
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(abstract_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create_abstract(21, "decimal", "%1.", "right", 540, 270)
    create_abstract(22, "bullet", "•", "left", 540, 270)
    create_abstract(23, "bullet", "☐", "left", 540, 270, "DejaVu Sans")


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def new_number_instance(doc, abstract_id):
    global NEXT_NUM_ID
    NEXT_NUM_ID += 1
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(NEXT_NUM_ID))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return NEXT_NUM_ID


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.font.size = Pt(11)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.line_spacing = 1.2


def add_paragraph_shading(paragraph, fill=LIGHT_BLUE, border=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_callout(doc, label, text, color=BLUE, fill=LIGHT_BLUE):
    paragraph = doc.add_paragraph(style="Callout")
    set_run(paragraph.add_run(f"{label}: "), bold=True, color=color)
    set_run(paragraph.add_run(text), color=NAVY)
    add_paragraph_shading(paragraph, fill=fill, border=color)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(paragraph.add_run(bold_prefix), bold=True, color=NAVY)
        set_run(paragraph.add_run(text[len(bold_prefix):]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_bullets(doc, items, checkbox=False):
    for item in items:
        paragraph = doc.add_paragraph()
        apply_num(paragraph, 23 if checkbox else 22)
        set_run(paragraph.add_run(item))
        set_keep(paragraph)


def add_steps(doc, items):
    num_id = new_number_instance(doc, 21)
    for title, detail in items:
        paragraph = doc.add_paragraph()
        apply_num(paragraph, num_id)
        set_run(paragraph.add_run(f"{title}. "), bold=True, color=NAVY)
        set_run(paragraph.add_run(detail))
        set_keep(paragraph)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("Saraiva.ai  |  Sites com ChatGPT  |  "), size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run(run, size=8.5, color=MUTED)


def add_page_title(doc, number, title, subtitle):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    set_run(kicker.add_run(f"MÓDULO {number}"), size=9, bold=True, color=BLUE)
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    set_run(heading.add_run(title), size=20, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    set_run(sub.add_run(subtitle), size=11.5, color=MUTED)


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_prompt_appendix(doc):
    lines = PROMPT_SOURCE.read_text(encoding="utf-8").splitlines()
    started = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line or line == "---":
            continue
        if line.startswith("# "):
            if started:
                add_page_break(doc)
            heading = doc.add_paragraph(style="Heading 1")
            heading.paragraph_format.space_before = Pt(0)
            heading.paragraph_format.space_after = Pt(8)
            set_run(heading.add_run(line[2:]), size=18, bold=True, color=NAVY)
            started = True
            continue
        if line.startswith("## "):
            if line == "## 8. Antes e depois ou portfólio":
                add_page_break(doc)
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(3)
            set_run(heading.add_run(line[3:]), size=11, bold=True, color=BLUE)
            set_keep(heading, keep_next=True)
            continue
        if line.startswith("* "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.14)
            paragraph.paragraph_format.space_after = Pt(1.5)
            set_run(paragraph.add_run(f"• {line[2:]}"), size=9.2, color=INK)
            set_keep(paragraph)
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.02
        if line == "`@Sites`":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(paragraph.add_run("@Sites"), size=15, bold=True, color=BLUE)
            add_paragraph_shading(paragraph, fill=LIGHT_BLUE, border=BLUE)
        else:
            set_run(paragraph.add_run(line.replace("`", "")), size=9.2, color=INK)
            set_keep(paragraph)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_numbering(doc)
    add_footer(doc.sections[0])

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(78)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("APOSTILA PRÁTICA • EDIÇÃO 1"), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("Sites com ChatGPT\nque vendem"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run(
        subtitle.add_run(
            "Do problema comercial ao site publicado:\n"
            "um método simples, prático e didático."
        ),
        size=15,
        color=MUTED,
    )
    cover_callout = doc.add_paragraph(style="Callout")
    cover_callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        cover_callout.add_run(
            "O ChatGPT acelera a execução. A estratégia determina se alguém vai pagar por ela."
        ),
        size=12,
        bold=True,
        color=NAVY,
    )
    add_paragraph_shading(cover_callout, fill=LIGHT_BLUE, border=BLUE)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(86)
    set_run(author.add_run("Saraiva.ai"), size=12, bold=True, color=BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(meta.add_run("Tecnologia para humanos • 2026"), size=9.5, color=MUTED)

    add_page_break(doc)
    add_page_title(doc, "01", "Comece pelo dinheiro, não pelo layout", "O site é uma peça do processo comercial.")
    add_body(
        doc,
        "Uma empresa não compra páginas, animações ou blocos de código. Ela compra um resultado: mais contatos, mais credibilidade, mais agendamentos, mais pedidos ou menos atrito para vender.",
    )
    add_callout(
        doc,
        "Regra central",
        "Antes de abrir o ChatGPT, escreva a ação que o visitante precisa realizar.",
    )
    add_steps(doc, [
        ("Escolha o negócio", "Trabalhe com um segmento que você consegue observar e abordar."),
        ("Localize o problema", "Descubra onde o processo comercial perde atenção, confiança ou conversão."),
        ("Defina a ação", "Escolha uma única conversão principal: chamar, agendar, pedir orçamento ou comprar."),
        ("Só então crie o site", "A estrutura e a tecnologia devem servir à ação escolhida."),
    ])
    add_body(doc, "Preencha agora:")
    add_bullets(doc, [
        "Negócio: ______________________________________________",
        "Problema comercial: _____________________________________",
        "Ação principal do visitante: ______________________________",
        "Resultado esperado: ______________________________________",
    ], checkbox=True)

    add_page_break(doc)
    add_page_title(doc, "02", "O briefing de 10 minutos", "Cinco respostas evitam horas construindo a coisa errada.")
    add_steps(doc, [
        ("Público", "Quem precisa reconhecer imediatamente que aquele site foi feito para ele?"),
        ("Problema valioso", "Qual problema essa pessoa já percebe e quer resolver?"),
        ("Oferta", "O que exatamente será apresentado ou vendido?"),
        ("Prova", "Que evidência real aumenta a confiança: processo, demonstração, portfólio ou dado verificável?"),
        ("CTA", "Qual é o próximo passo mais simples que pode ser medido?"),
    ])
    add_callout(
        doc,
        "Teste de clareza",
        "Se você não consegue explicar a oferta em uma frase, o site ainda não está pronto para ser criado.",
        color=GOLD,
        fill="FFF7E8",
    )
    add_body(doc, "Frase de oferta:")
    add_body(
        doc,
        "Eu ajudo [PÚBLICO] a sair de [PROBLEMA] e chegar a [RESULTADO] por meio de [MECANISMO], com [PRÓXIMO PASSO].",
    )

    add_page_break(doc)
    add_page_title(doc, "03", "O prompt mestre", "Use o modo Work do ChatGPT com @Sites.")
    add_callout(
        doc,
        "Configuração correta",
        "Abra uma nova conversa, selecione o modo Work, ative ou mencione @Sites e cole o prompt completo do apêndice.",
        color=BLUE,
        fill=LIGHT_BLUE,
    )
    add_steps(doc, [
        ("Abra o modo Work", "Use uma conversa dedicada para o projeto e mantenha nela todo o contexto do negócio."),
        ("Acione @Sites", "Informe explicitamente que o resultado precisa ser um site navegável, funcional e pronto para publicação."),
        ("Cole o prompt completo", "Não remova as regras de veracidade, conversão, SEO, desempenho ou acessibilidade."),
        ("Acompanhe até funcionar", "Não encerre na estratégia: revise as páginas, os botões, o formulário e a versão mobile."),
    ])
    add_body(
        doc,
        "O prompt completo para a Scarlett Makeup está no apêndice. Ele pode ser adaptado para outro negócio trocando somente os dados factuais e os campos marcados como pendentes.",
    )
    add_bullets(doc, [
        "Substitua todos os campos [PREENCHER: ...] antes da publicação.",
        "Rejeite qualquer prova, depoimento, horário ou número inventado.",
        "Teste WhatsApp, ligação, formulário, páginas internas e versão mobile.",
        "Continue a execução até o site estar navegável e funcional.",
    ], checkbox=True)

    add_page_break(doc)
    add_page_title(doc, "04", "A estrutura que conduz à ação", "Cada seção responde uma dúvida e aproxima o visitante do CTA.")
    add_steps(doc, [
        ("Hero", "Diga para quem é, qual resultado oferece e qual ação deve ser tomada."),
        ("Problema", "Mostre que você entende a situação atual sem exagero ou manipulação."),
        ("Mecanismo", "Explique como a solução funciona em linguagem simples."),
        ("Benefícios", "Traduza recursos em mudanças percebidas pelo cliente."),
        ("Prova", "Use apenas evidência verificável e relevante para a decisão."),
        ("Oferta", "Mostre o que está incluído, condições e limites."),
        ("Objeções", "Responda às dúvidas que realmente impedem o próximo passo."),
        ("CTA final", "Repita a ação principal com contexto e baixo atrito."),
    ])
    add_callout(
        doc,
        "Corte",
        "Se uma seção não aumenta clareza, confiança ou ação, ela provavelmente não precisa existir.",
    )

    add_page_break(doc)
    add_page_title(doc, "05", "Copy que explica e move", "Texto bonito não basta; cada bloco precisa cumprir uma função.")
    add_body(doc, "Modelo de hero:")
    add_callout(
        doc,
        "Promessa",
        "[RESULTADO] para [PÚBLICO] por meio de [MECANISMO], sem [ATRITO QUE PODE SER REDUZIDO].",
    )
    add_body(doc, "Modelo de subtítulo:")
    add_callout(
        doc,
        "Contexto",
        "Explique o que acontece, em quanto tempo aparece a primeira evidência e qual é o próximo passo.",
        color=GREEN,
        fill="EAF8F4",
    )
    add_body(doc, "Checklist de copy:")
    add_bullets(doc, [
        "A primeira tela deixa claro para quem é.",
        "O resultado é específico, reconhecível e plausível.",
        "O mecanismo é explicado sem jargão.",
        "A prova é real e está perto da afirmação que sustenta.",
        "O CTA usa um verbo e diz o que acontece depois.",
        "Não há urgência, garantia ou escassez inventada.",
    ], checkbox=True)

    add_page_break(doc)
    add_page_title(doc, "06", "Design mobile-first", "A maioria dos visitantes decidirá em uma tela pequena.")
    add_bullets(doc, [
        "Uma ideia principal por bloco.",
        "Título curto, contraste forte e corpo legível.",
        "Botão principal visível sem disputar atenção com vários links.",
        "Espaçamento suficiente para leitura e toque.",
        "Imagens úteis, leves e relacionadas à oferta.",
        "Formulário com o mínimo de campos necessário.",
        "Estados claros de carregamento, erro e sucesso.",
    ], checkbox=True)
    add_callout(
        doc,
        "Teste de 5 segundos",
        "Mostre a primeira tela a alguém. Pergunte: para quem é, o que oferece e onde clicar. Se a resposta não vier rápido, simplifique.",
        color=GOLD,
        fill="FFF7E8",
    )
    add_body(
        doc,
        "Evite usar animação para compensar uma mensagem fraca. Movimento deve orientar atenção, confirmar uma ação ou explicar uma transição.",
    )

    add_page_break(doc)
    add_page_title(doc, "07", "Do ChatGPT ao site publicado", "Construa em ciclos curtos e valide cada parte.")
    add_steps(doc, [
        ("Estratégia", "Aprove oferta, público, problema e CTA."),
        ("Wireframe", "Aprove ordem das seções e hierarquia da página."),
        ("Copy", "Revise fatos, promessa, prova e chamadas para ação."),
        ("Visual", "Defina cores, tipografia, imagens e estados mobile."),
        ("Código", "Peça componentes simples, acessíveis e responsivos."),
        ("Integrações", "Conecte formulário, WhatsApp, agenda, CRM e métricas."),
        ("Publicação", "Configure domínio, HTTPS, metadados e analytics."),
        ("Teste", "Percorra o caminho como visitante antes de divulgar."),
    ])
    add_body(doc, "Prompt de correção:")
    add_callout(
        doc,
        "Iteração",
        "Analise esta página como especialista em conversão e UX. Liste apenas os três problemas que mais prejudicam clareza, confiança ou ação. Para cada problema, explique a evidência e proponha uma correção concreta. Não reescreva tudo.",
    )

    add_page_break(doc)
    add_page_title(doc, "08", "Integrações mínimas", "O site precisa terminar em uma operação que responde.")
    add_steps(doc, [
        ("Entrada", "Formulário, WhatsApp, agenda ou checkout."),
        ("Registro", "Salve origem, campanha, página e mensagem do lead."),
        ("Resposta", "Confirme imediatamente o que aconteceu e o próximo passo."),
        ("Responsável", "Defina quem assume quando o lead pede ajuda ou compra."),
        ("Follow-up", "Crie lembretes para conversas sem resposta."),
        ("Medição", "Acompanhe visita, clique, lead, conversa qualificada e venda."),
    ])
    add_callout(
        doc,
        "Atenção",
        "Automação não é licença para spam. Use contexto, permissão, opt-out e limites de frequência.",
        color=RED,
        fill="FFF0F0",
    )

    add_page_break(doc)
    add_page_title(doc, "09", "Checklist de publicação", "Não divulgue antes de testar o caminho inteiro.")
    add_bullets(doc, [
        "Domínio e HTTPS funcionando.",
        "Título e descrição corretos ao compartilhar o link.",
        "Página responsiva em celular real.",
        "Botões e links apontando para o destino correto.",
        "Formulário validado e confirmação exibida.",
        "Lead registrado no CRM ou planilha.",
        "Notificação entregue ao responsável.",
        "Analytics recebendo eventos.",
        "Política de privacidade e informações de contato presentes.",
        "Sem placeholders, provas falsas ou textos genéricos.",
        "Velocidade aceitável em rede móvel.",
        "Compra ou agendamento testado de ponta a ponta.",
    ], checkbox=True)
    add_body(doc, "Critério de aceite:")
    add_callout(
        doc,
        "Pronto",
        "Uma pessoa consegue entender a oferta, realizar a ação e receber confirmação sem depender de você explicar o caminho.",
        color=GREEN,
        fill="EAF8F4",
    )

    add_page_break(doc)
    add_page_title(doc, "10", "Como vender o serviço", "Venda a melhoria comercial, não a quantidade de páginas.")
    add_steps(doc, [
        ("Diagnóstico", "Mostre onde o processo atual perde atenção, confiança ou resposta."),
        ("Escopo", "Defina a ação principal e o que será integrado."),
        ("Protótipo", "Apresente a primeira tela e a jornada antes do site completo."),
        ("Construção", "Produza em ciclos curtos com aprovações objetivas."),
        ("Entrega", "Publique, documente acessos e ensine a operação."),
        ("Evolução", "Revise dados reais e priorize a próxima melhoria."),
    ])
    add_body(doc, "Pergunta comercial de abertura:")
    add_callout(
        doc,
        "Conversa",
        "Hoje o seu site perde mais oportunidades por falta de clareza, confiança, resposta rápida ou acompanhamento do lead?",
    )
    add_body(
        doc,
        "Não prometa faturamento. Combine o que será entregue, como será medido e quais resultados dependem de tráfego, oferta, atendimento e execução do cliente.",
    )

    add_page_break(doc)
    add_page_title(doc, "11", "Painel de melhoria", "Meça o funil completo, não apenas visitas.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2500, 3430, 3430])
    headers = ["Etapa", "Métrica", "Pergunta"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        set_run(paragraph.add_run(value), size=10, bold=True, color=NAVY)
    rows = [
        ("Atenção", "Visitas e origem", "De onde chegam as pessoas certas?"),
        ("Interesse", "Cliques no CTA", "A promessa cria vontade de avançar?"),
        ("Lead", "Formulários ou conversas", "O próximo passo tem atrito?"),
        ("Qualificação", "Leads com problema real", "O público alcançado tem aderência?"),
        ("Venda", "Pagamentos confirmados", "A oferta e o atendimento fecham o ciclo?"),
        ("Aprendizado", "Motivos de perda", "Qual bloqueio aparece com mais frequência?"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            set_run(paragraph.add_run(value), size=9.5, bold=index == 0, color=NAVY if index == 0 else INK)
    set_table_geometry(table, [2500, 3430, 3430])
    add_body(
        doc,
        "Registre uma linha de base antes de mudar a página. Depois compare a mesma janela, origem e oferta. Um caso isolado é sinal; repetição em condições semelhantes vira aprendizado.",
    )

    add_page_break(doc)
    add_page_title(doc, "12", "Próximo movimento", "Você pode aprender o processo ou receber a automação montada.")
    add_body(
        doc,
        "Use esta apostila para construir o primeiro site completo. O objetivo é chegar a uma prova real: página publicada, ação funcionando e dados chegando.",
    )
    add_callout(
        doc,
        "Rota DIY",
        "Preencha o briefing, use o prompt mestre, construa em etapas e valide com o checklist.",
        color=GREEN,
        fill="EAF8F4",
    )
    add_callout(
        doc,
        "Rota pronta",
        "Se quiser briefing, geração, integrações, publicação, CRM, pagamento e follow-up no mesmo fluxo, responda PRONTA na conversa do Instagram.",
        color=BLUE,
        fill=LIGHT_BLUE,
    )
    add_body(doc, "Seu compromisso de 48 horas:")
    add_bullets(doc, [
        "Escolher um negócio e um problema comercial.",
        "Preencher o briefing de cinco respostas.",
        "Gerar e revisar a primeira estrutura.",
        "Publicar uma versão simples com um CTA funcionando.",
        "Registrar a primeira evidência real.",
    ], checkbox=True)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(24)
    set_run(closing.add_run("Menos efeito. Mais clareza. Mais ação."), size=15, bold=True, color=NAVY)
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(brand.add_run("Saraiva.ai • Tecnologia para humanos"), size=10, color=BLUE, bold=True)

    add_page_break(doc)
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(8)
    set_run(intro.add_run("APÊNDICE • COPIE A PARTIR DA PRÓXIMA LINHA"), size=10, bold=True, color=BLUE)
    appendix_title = doc.add_paragraph()
    appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    appendix_title.paragraph_format.space_after = Pt(12)
    set_run(appendix_title.add_run("Prompt completo\nModo Work + @Sites"), size=26, bold=True, color=NAVY)
    add_callout(
        doc,
        "Antes de colar",
        "Revise os dados do negócio. O prompt proíbe a invenção de avaliações, horários, garantias e resultados.",
        color=GREEN,
        fill="EAF8F4",
    )
    add_page_break(doc)
    add_prompt_appendix(doc)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
